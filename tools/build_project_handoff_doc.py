from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'darktide-survivors-project-handoff.docx'

NAVY = '172333'
BLUE = '2E74B5'
GOLD = 'B28A47'
LIGHT = 'E8EEF5'
MUTED = '5B6573'
RED = '9B1C1C'
GREEN = '1F6B43'

def set_font(run, name='Calibri', size=11, color='1F2933', bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for side, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+side))
        if node is None: node = OxmlElement('w:'+side); tcMar.append(node)
        node.set(qn('w:w'), str(val)); node.set(qn('w:type'), 'dxa')

def set_cell_width(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(dxa)); tcW.set(qn('w:type'), 'dxa')

def table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'), 'dxa')
    ind = tblPr.find(qn('w:tblInd'))
    if ind is None: ind = OxmlElement('w:tblInd'); tblPr.append(ind)
    ind.set(qn('w:w'), '120'); ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i]); cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def mark_header_row(row):
    """Mark a repeated table header so screen readers and Word understand it."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def add_para(doc, text='', style=None, size=11, color='1F2933', bold=False, italic=False, before=0, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.25
    if align is not None: p.alignment = align
    r = p.add_run(text); set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); set_font(r, size=10.5)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text); set_font(r, size=10.5); return p

def add_callout(doc, label, text, fill='F4F6F9', accent=GOLD):
    t = doc.add_table(rows=1, cols=1); table_geometry(t, [9360]); c=t.cell(0,0); shade(c, fill); cell_margins(c, 130, 180, 130, 180)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.2
    r=p.add_run(label+'  '); set_font(r, size=10.5, color=accent, bold=True)
    r=p.add_run(text); set_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def add_status_table(doc, rows):
    table=doc.add_table(rows=1, cols=4); table_geometry(table, [2100, 3000, 2500, 1760])
    mark_header_row(table.rows[0])
    headers=['模块','已完成','关键入口','状态']
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; shade(c, NAVY); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(h); set_font(r,size=10,color='FFFFFF',bold=True)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            shade(cells[i], 'F7F9FB' if len(table.rows)%2==0 else 'FFFFFF')
            p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(val); set_font(r,size=9.2,color=GREEN if i==3 and val.startswith('已') else '1F2933',bold=(i==3))
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.paragraph_format.keep_with_next=True
    r=p.add_run(text); set_font(r,size={1:16,2:13,3:11.5}[level],color=BLUE if level<3 else '1F4D78',bold=True)
    return p

def configure(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1); sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)
    normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for level, size, color, before, after in [(1,16,BLUE,18,10),(2,13,BLUE,14,7),(3,12,'1F4D78',10,5)]:
        st=doc.styles[f'Heading {level}']; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    for name in ['List Bullet','List Bullet 2','List Number']:
        st=doc.styles[name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); st.font.size=Pt(10.5); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.25
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run('DarkEscaper · 项目改动交接文档'); set_font(r,size=8,color=MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('供后续模型继续开发使用 · 生成于 2026-08-10'); set_font(r,size=8,color=MUTED)

def build():
    doc=Document(); configure(doc)
    add_para(doc,'DARKESCAPER',size=10,color=GOLD,bold=True,after=3,align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,'项目改动全量交接文档',size=26,color=NAVY,bold=True,after=5,align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,'美术资源 · 动画系统 · 地形 · UI · 性能 · 发布状态',size=13,color=MUTED,after=22,align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc,'交接目的','本文件汇总当前对话中已经落地到仓库的项目改动，帮助其他模型快速理解现状、定位代码、复现实验并继续优化。文档同时明确哪些问题已修复、哪些入口仍值得重点复核。',fill='FFF8E8')
    add_status_table(doc,[
        ('仓库','darktide-survivors','main · ad1bc23','已同步'),
        ('生产站点','darktide-survivors.pages.dev','Cloudflare Pages','已部署'),
        ('图集','assets/sprites/atlas.*','444 项 / 2527 帧','已通过'),
        ('视觉风格','写实复古中世纪像素','黑铁/黄铜/哥特 UI','统一基线'),
        ('运行性能','400 敌人同屏','约 34–41 FPS 测试均值','已通过'),
    ])
    heading(doc,'一、项目背景与最终美术方向',1)
    add_para(doc,'本项目是一个离线优先、支持触屏和联机快照的中世纪暗黑生存游戏。美术方向经过多次讨论后确定为：偏写实、细腻、复古像素、略带恐怖但不走纯血腥惊悚；角色与怪物要有岁月感和区分度，UI 采用哥特/巴洛克古典视觉。')
    add_bullet(doc,'像素规则：硬 Alpha（0/255）、最近邻缩放、禁止普通图片式抗锯齿边缘；素材制作阶段可用洋红色键背景，本地统一去背。')
    add_bullet(doc,'光源与色彩：深色地面、金铜 UI、黑铁描边、红色危险提示、蓝白寒霜、青蓝电弧、腐绿沼泽；素材必须保持强轮廓和可辨识剪影。')
    add_bullet(doc,'动作规则：角色/敌人采用独立帧条；角色至少四方向（down/right/left/up），动作包括 idle、walk、attack、hurt/death；帧间脚底锚点稳定。')
    add_bullet(doc,'运行接口：保留 SpriteGen.get() / SpriteGen.frames()，优先读取 atlas，资源缺失时走程序生成 fallback。')
    heading(doc,'二、已落地改动总览',1)
    add_status_table(doc,[
        ('角色动作','四方向、待机/跑步/攻击/受击/死亡，修复残帧与反向','tools/art/repair-sprite-grids.py；register-art-v2.js','已完成'),
        ('敌人/Boss','普通、精英、Boss 独立动作与尺寸，巨腐史莱姆放大','js/entities.js；js/config.js','已完成'),
        ('武器特效','贤者光弹、圣女光环、寒霜、Tesla、秘典恶魔、火池','js/weapons.js','已完成'),
        ('地形','地面/道路/草地/沼泽纹理；大型固定泥边水坑','js/main.js；js/engine.js；art-v4','已完成'),
        ('装饰排序','固定世界位置、低矮层/前景层、按 y 排序、扩大间距','js/main.js；js/engine.js','已完成'),
        ('UI','巴洛克边框、古典纹理、HUD/Boss/百科/索敌布局','css/art-v4.css；js/ui.js；js/minimap.js','已完成'),
        ('触屏/性能','双指触控、400 敌人性能、预加载图集','js/ui.js；测试脚本','已完成'),
    ])
    heading(doc,'三、按用户需求逐条映射',1)
    items=[
        ('1. 贤者光弹闪烁','p_bolt 在绘制时固定使用首帧，避免旧动作条空白帧循环导致闪烁；注册入口为 tools/art/register-art-v2.js 的 p_bolt。继续优化时应优先替换 source，而不是恢复随机帧。','js/weapons.js：p_bolt 绘制分支；assets/art-v2/sprites/vfx/projectiles_actions.png'),
        ('2. 沼泽固定、多形、大尺寸','旧四格源图存在跨格裁切；当前改为独立的 swamp_puddle_large_master.png，生成四个硬像素镜像/旋转变体，运行时 320×224 逻辑帧，地图绘制约 248–364×150–218 像素；世界坐标由哈希固定，数量阈值约 13%。','assets/art-v4/sources/terrain/swamp_puddle_large_master.png；tools/art/build-art-v4.py；js/main.js；js/engine.js'),
        ('3. Boss 出场框','hud-warn 设置 max-width、min-width、padding、双线金框；boss-wrap 使用 calc(100vw - 34px) 并固定 box-sizing，文字不会溢出。','css/art-v4.css；js/ui.js'),
        ('4. 光环/寒霜','光环固定完整环形帧，只做低幅度 alpha 呼吸；寒霜改用 vfx_frost_radial，按半径进度由中心向 360° 外扩。','js/weapons.js；assets/art-v4/sprites/vfx/frost_radial_actions.png'),
        ('5. 所罗门秘典','满级进化由书本弹幕改为 demon 子弹实体，使用 vfx_spirit，环绕玩家并在 280 范围主动寻敌，持续命中。','js/weapons.js；js/config.js'),
        ('6. Tesla 电塔','塔身固定使用 tesla_tower[0]，只在出生前 0.42 秒进行 y 方向钻出；顶部电球/电弧独立叠加，避免换帧抽搐。','js/weapons.js'),
        ('7. 巨腐史莱姆','CFG.ENEMIES.slime_big 的碰撞半径由 15 调到 21，绘制缩放区分普通腐液史莱姆。','js/config.js；js/entities.js；register-art-v2.js'),
        ('8. 怪物/Boss 弹幕','ImageGen 生成毒液、奥术、血球、骨矢、地狱火、冰针、电球、邪眼八类静态弹幕；实体根据颜色选择素材。','assets/art-v4/sources/vfx/frost_enemy_projectiles_master.png；js/entities.js'),
        ('9. 磁铁','V4 pickup sheet 采用红/蓝两侧；程序 fallback magnet 也明确 red/blue。若继续验收，应检查线上 atlas 的 magnet 预览。','assets/art-v4/sprites/pickups/pickup_gems.png；js/sprites.js'),
        ('10. 道具大小','掉落物 itemScale 已从放大版回调到 chest 0.58、coin 0.47、普通道具 0.55。','js/entities.js'),
        ('11. 商人','商人三种动作注册 renderScale 0.40；普通待机固定首帧消除闪烁；商人箭矢起点改到适配新比例。','js/merchant.js；register-art-v2.js'),
        ('12. 百科图标','百科头部、Boss 头部、进化图标 CSS 尺寸提高，并调整卡片/背景布局。','js/ui.js；css/art-v4.css'),
        ('13. 全 UI 巴洛克','新增 baroque_frame 图像边框，主菜单、暂停、百科、HUD 使用黑铁/黄铜/古典纹理统一样式。','css/art-v4.css；assets/art-v4/sprites/ui/baroque_menu_frame.png'),
        ('14. 左右方向','方向表修正为 down=0、right=1、left=2、up=3；绘制时不再二次翻转。当前骑士预览已确认左右朝向正确。','tools/art/register-art-v2.js；js/entities.js'),
        ('15. 左上 HUD','移除“生命·装备”标题；hp-wrap 采用金铜外框、黑色描边、内嵌层次。','js/ui.js；css/art-v4.css'),
        ('16. 索敌文案','小地图下方文案改成“索敌方式：最近敌人/最低血量/最高血量”。','js/minimap.js；js/weapons.js'),
        ('17. 角色尺寸/缺头','repair-sprite-grids.py 对四方向动作全部启用严格修复，按完整轮廓替换残帧；重复帧增加 1 像素硬像素步态相位；脚底基线统一。','tools/art/repair-sprite-grids.py；assets/art-v4/repaired/characters/'),
        ('18. 石头地板图层','低矮装饰（岩石、根须、骨头等不在 isTallDecor 中）通过 pass=ground 在角色前后层之前绘制；高大装饰再按 back/front 处理。','js/main.js：drawDecor 调用顺序与 isTallDecor'),
        ('19. 装饰互不遮挡','DECOR_CELL 调到 304；位置固定在格心附近的小范围偏移；同一 pass 的装饰按 y、x 排序，避免遍历顺序引起相互覆盖跳层。','js/engine.js：decorEntry/forEachDecor；js/main.js：drawDecor'),
    ]
    for title, body, paths in items:
        heading(doc,title,2); add_para(doc,body); add_para(doc,'关键入口：'+paths,size=9.5,color=MUTED,italic=True,after=8)
    heading(doc,'四、资源生产与图集管线',1)
    add_number(doc,'ImageGen 负责设计稿/动作源图。透明素材使用纯 #ff00ff 背景生成，避免直接把洋红背景交给运行时。')
    add_number(doc,'build-art-v4.py 执行去背、硬 Alpha、调色板量化、最近邻缩放，并输出 assets/art-v4/sprites。当前脚本支持 --skip-repair，仅重制 V4 美术而不重新处理角色修复。')
    add_number(doc,'register-art-v2.js 将源图注册到 tools/art/art-manifest.json；atlas 构建输出 assets/sprites/atlas.png、atlas.json、atlas-data.js 和 previews。')
    add_number(doc,'build-atlas.js 生成并验证图集；最近一次结果为 444 项、2527 帧，ANCHOR/ATLAS/ANIM/ART 均通过。')
    add_callout(doc,'注意','不要直接把 ImageGen 生成的网格图当作严格网格。若模型跨格，必须改成独立源图或先做 alpha 连通域切分，否则会再次出现半张水坑、重复角色、闪烁。',fill='FFF8E8',accent=RED)
    heading(doc,'五、关键文件地图',1)
    files=[
        ('js/main.js','主循环、地面/沼泽绘制、装饰 pass、菜单背景、相机坐标。'),
        ('js/engine.js','输入、地形效果、沼泽/道路碰撞、固定装饰生成与碰撞。'),
        ('js/entities.js','玩家/敌人/掉落物/敌方弹幕绘制和动作选择。'),
        ('js/weapons.js','全部武器发射、特效、光环、寒霜、Tesla、秘典恶魔、火池。'),
        ('js/merchant.js','商人购买、装死、弓箭攻击和动作状态。'),
        ('js/ui.js','主菜单、HUD、Boss 条、暂停、百科、触屏按钮。'),
        ('js/minimap.js','小地图、地形标记和索敌方式文案。'),
        ('css/art-v4.css','当前最后加载的巴洛克/黑铁黄铜视觉覆盖层。'),
        ('tools/art/repair-sprite-grids.py','修复跨格、残帧、重复帧和四方向角色动作。'),
        ('tools/art/build-art-v4.py','ImageGen 源图到硬像素 V4 资产。'),
        ('tools/art/register-art-v2.js','V4 图集注册表和逻辑尺寸/锚点。'),
        ('test/art-probe.js','图集锚点、帧差异、资源结构与浏览器加载验收。'),
        ('test/headless.js','开局、暂停、武器进化、联机、性能和完整冒烟。'),
    ]
    add_status_table(doc,[(a,b,'—','已知入口') for a,b in files])
    heading(doc,'六、验证记录',1)
    add_bullet(doc,'图集：444 assets / 2527 frames；ANCHOR OK、ATLAS OK、ANIM OK、ART OK。')
    add_bullet(doc,'性能：400 敌人同屏测试约 34–41 FPS，p95 约 33.4ms。')
    add_bullet(doc,'地形：TERRAIN OK（道路 ×1.2、沼泽 ×0.6、装饰碰撞）。')
    add_bullet(doc,'触屏：TOUCH OK，移动时仍可用第二指暂停/切换索敌。')
    add_bullet(doc,'完整冒烟：12 把武器满级/进化、合作模式快照、死亡/复活、商人、掉落物均通过。')
    add_callout(doc,'已知测试噪声','无头测试环境有时会报告 vfx_frost_radial 未知素材，因为它使用旧版 fallback 初始化；真实浏览器 atlas-probe 已确认该资源存在并通过加载验收。继续开发时应把测试 stub 也登记该别名，消除噪声。',fill='FFF8E8',accent=RED)
    heading(doc,'七、Git 与部署',1)
    add_para(doc,'当前 main 分支最近提交：')
    for c in ['ad1bc23 fix: stabilize decor layering and placement','66c6423 fix: rebuild large fixed swamp and firepool loop','438cc9c fix: stabilize combat art terrain and hud','f9b99c5 feat: 重制地形特效动作与触屏界面']:
        add_bullet(doc,c)
    add_para(doc,'已推送到 GitHub 与 Gitee。GitHub Actions workflow deploy.yml 会自动执行 Cloudflare Pages 部署；最近一次部署成功。生产地址：https://darktide-survivors.pages.dev/')
    heading(doc,'八、给后续模型的继续开发建议',1)
    for s in [
        '先解决测试 stub 的 vfx_frost_radial 别名，再扩展任何新素材；保持真实 atlas 与 fallback 两套接口一致。',
        '若继续生成角色动作，优先生成单角色/单动作/单方向源图，再由本地脚本拼接，不要让模型一次生成 4×8 严格网格。',
        '每次改动地图装饰，都要同时检查 decorEntry 的世界坐标、drawDecor 的 pass 顺序和 resolveDecorCollision 的半径。',
        '每次改动 UI，至少检查主菜单、进入地图 HUD、Boss 出场、暂停、百科、升级、商店、联机 HUD 七个界面。',
        '每次素材/渲染改动都运行 node test/art-probe.js、node test/terrain-touch-probe.js、node test/headless.js，并在浏览器中实际看一遍。',
        '提交前保留可回退 commit；不要覆盖原始 ImageGen source，新增版本化文件并在 manifest 中登记。',
    ]: add_bullet(doc,s)
    heading(doc,'九、交接完成判定清单',1)
    for s in ['[ ] 真实浏览器中确认贤者光弹连续飞行 5 秒无闪烁','[ ] 真实浏览器中确认大型沼泽随相机移动保持世界固定','[ ] 确认火池只播放纯燃烧帧且火焰不消失','[ ] 检查四向角色的左右朝向、上方向头部和脚底高度','[ ] 检查所有低矮石头/根须在角色下方、高大树木按 y 正确遮挡','[ ] 检查两个相邻装饰不会互相穿插','[ ] 检查 Boss 提示、百科、移动端 HUD 在窄屏不溢出','[ ] 测试通过后提交并等待 Cloudflare workflow 成功']: add_bullet(doc,s)
    doc.core_properties.title='DarkEscaper 项目改动全量交接文档'; doc.core_properties.subject='项目美术、动画、地形、UI、测试和发布状态'; doc.core_properties.author='Codex'
    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT); print(OUT)

if __name__=='__main__': build()
